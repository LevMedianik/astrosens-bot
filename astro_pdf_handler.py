import os
from pdfminer.high_level import extract_text
import docx
from uuid import uuid4
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_huggingface import HuggingFaceEmbeddings
from langchain.docstore.document import Document
from dotenv import load_dotenv
import requests

load_dotenv()

# Директории
DATA_DIR = "./data"
INDEX_DIR = "./faiss_index"
os.makedirs(DATA_DIR, exist_ok=True)
os.makedirs(INDEX_DIR, exist_ok=True)

# Сохраняем загружаемый файл
def save_file(file_bytes, original_filename):
    file_id = str(uuid4())
    filename = f"{file_id}_{original_filename}"
    filepath = os.path.join(DATA_DIR, filename)
    with open(filepath, 'wb') as f:
        f.write(file_bytes)
    return filepath

# Извлечение текста из PDF
def extract_text_from_pdf(filepath):
    return extract_text(filepath)

# Извлечение текста из DOCX
def extract_text_from_docx(filepath):
    doc = docx.Document(filepath)
    return '\n'.join([para.text for para in doc.paragraphs])

# Извлечение текста из TXT
def extract_text_from_txt(filepath):
    with open(filepath, 'r', encoding='utf-8') as f:
        return f.read()

# Универсальная функция извлечения
def extract_text_from_file(filepath):
    if filepath.endswith(".pdf"):
        return extract_text_from_pdf(filepath)
    elif filepath.endswith(".docx"):
        return extract_text_from_docx(filepath)
    elif filepath.endswith(".txt"):
        return extract_text_from_txt(filepath)
    else:
        raise ValueError("❌ Поддерживаются только PDF, DOCX и TXT файлы.")

# Индексация текста
def index_text_with_faiss(text):
    splitter = RecursiveCharacterTextSplitter(chunk_size=300, chunk_overlap=50)
    chunks = splitter.split_text(text)
    documents = [Document(page_content=chunk) for chunk in chunks]

    embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
    vectorstore = FAISS.from_documents(documents, embedding=embeddings)
    vectorstore.save_local(INDEX_DIR)
    return vectorstore

# Загрузка индекса
def load_existing_index():
    index_file = os.path.join(INDEX_DIR, "index.faiss")
    if not os.path.exists(index_file):
        return None
    embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
    return FAISS.load_local(INDEX_DIR, embeddings, allow_dangerous_deserialization=True)

# Настройки LLM
OPENROUTER_API_KEY = os.getenv("OPENROUTER_API_KEY")
HEADERS = {
    "Authorization": f"Bearer {OPENROUTER_API_KEY}",
    "HTTP-Referer": "http://localhost",
    "X-Title": "AstroSens"
}
LLM_MODEL = "deepseek/deepseek-chat-v3-0324:free"

# Запрос к модели
def ask_llm(prompt: str):
    body = {
        "model": LLM_MODEL,
        "messages": [
            {"role": "system", "content": "Ты — научный ассистент, отвечающий на вопросы по астрономии и астрофизике."},
            {"role": "user", "content": prompt}
        ],
        "temperature": 0.3
    }
    response = requests.post("https://openrouter.ai/api/v1/chat/completions", headers=HEADERS, json=body)
    if response.status_code != 200:
        return f"Ошибка запроса к LLM: {response.status_code} - {response.text}"
    return response.json()['choices'][0]['message']['content']

# Вопрос к индексу
def query_index(question: str, announce: bool = False):
    vectorstore = load_existing_index()
    if not vectorstore:
        return "❌ База знаний не найдена. Пожалуйста, загрузите документ."

    retriever = vectorstore.as_retriever(search_type="mmr", search_kwargs={"k": 4})
    results = retriever.get_relevant_documents(question)
    context = "\n".join([doc.page_content.strip() for doc in results])
    full_prompt = f"Контекст:\n{context}\n---\nВопрос: {question}\nОтвет:"
    reply = ask_llm(full_prompt)
    return ("🔍 Ищу ответ...", reply) if announce else reply

# Суммаризация
def summarize_pdf(announce: bool = False):
    vectorstore = load_existing_index()
    if not vectorstore:
        return "Индекс не найден. Пожалуйста, загрузите PDF-документ."

    retriever = vectorstore.as_retriever(search_type="mmr", search_kwargs={"k": 5})
    results = retriever.get_relevant_documents("Основное содержание файла")
    text = "\n".join([doc.page_content.strip() for doc in results])
    prompt = f"Ты — нейросеть, кратко пересказывающая статью. Перескажи это понятно и точно:\n{text}"
    reply = ask_llm(prompt)
    return ("📖 Пересказываю текст...", reply) if announce else reply
